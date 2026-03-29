"""Parse a docx CV and detect sections by style, formatting, and content signals."""

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.shared import Pt


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

@dataclass
class DocxSection:
    """A detected section in a docx CV."""
    name: str           # e.g. "Education", "Skills"
    semantic_type: str   # normalized: "education", "experience", "skills", "summary", "custom"
    header_para: int     # paragraph index of the header (-1 for implicit sections)
    start_para: int      # first content paragraph (after header)
    end_para: int        # last paragraph (inclusive) before next section


@dataclass
class DocxSectionMap:
    """Complete section map of a parsed docx CV."""
    header: Optional[DocxSection]    # name + title block at top
    contact: Optional[DocxSection]   # contact info block
    sections: List[DocxSection] = field(default_factory=list)
    unmapped_paras: List[int] = field(default_factory=list)

    def summary(self) -> str:
        """Human-readable summary of detected sections."""
        lines = []
        if self.header:
            lines.append(f"  [header]  — paragraphs {self.header.start_para}-{self.header.end_para}")
        if self.contact:
            lines.append(f"  [contact] — paragraphs {self.contact.start_para}-{self.contact.end_para}")
        for s in self.sections:
            lines.append(
                f"  [{s.semantic_type}] \"{s.name}\" "
                f"— paragraphs {s.start_para}-{s.end_para}"
            )
        if self.unmapped_paras:
            lines.append(f"  [unmapped] — {len(self.unmapped_paras)} paragraphs")
        return "\n".join(lines)


# ---------------------------------------------------------------------------
# Known section names → semantic type
# ---------------------------------------------------------------------------

_SECTION_MAP = {
    # summary
    "summary": "summary", "profile": "summary", "about": "summary",
    "about me": "summary", "objective": "summary", "professional summary": "summary",
    "career objective": "summary",
    # education
    "education": "education", "academic background": "education",
    "academic education": "education",
    # experience
    "experience": "experience", "work experience": "experience",
    "employment": "experience", "work history": "experience",
    "professional experience": "experience",
    # skills
    "skills": "skills", "technical skills": "skills",
    "core competencies": "skills", "competencies": "skills",
    "key skills": "skills",
    # projects
    "projects": "projects", "personal projects": "projects",
    "academic projects": "projects",
    # languages
    "languages": "languages",
    # military
    "military": "military", "military service": "military",
    "national service": "military",
    # volunteering
    "volunteering": "volunteering", "volunteer experience": "volunteering",
    "volunteer work": "volunteering", "community involvement": "volunteering",
    # certifications
    "certifications": "certifications", "certificates": "certifications",
    "licenses": "certifications", "licenses & certifications": "certifications",
    # publications
    "publications": "publications",
    # awards
    "awards": "awards", "honors": "awards", "achievements": "awards",
    "honors & awards": "awards",
    # references
    "references": "references",
}

# Set of all known names for quick lookup (Signal 3)
_KNOWN_NAMES = set(_SECTION_MAP.keys())

# Contact info regex patterns
_EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
_PHONE_RE = re.compile(r"[\d\-+()\s]{7,}")
_LINKEDIN_RE = re.compile(r"linkedin\.com", re.IGNORECASE)


# ---------------------------------------------------------------------------
# Section classification
# ---------------------------------------------------------------------------

def classify_section(header_text: str) -> str:
    """Map detected header text to a semantic type."""
    normalized = header_text.strip().lower()
    # Direct match
    if normalized in _SECTION_MAP:
        return _SECTION_MAP[normalized]
    # Partial match — check if any known name is contained in the text
    for name, sem_type in _SECTION_MAP.items():
        if name in normalized:
            return sem_type
    return "custom"


# ---------------------------------------------------------------------------
# Header detection signals
# ---------------------------------------------------------------------------

def _compute_median_font_size(paragraphs) -> Optional[int]:
    """Compute the median font size across all runs with explicit size."""
    sizes = []
    for para in paragraphs:
        for run in para.runs:
            if run.font.size:
                sizes.append(run.font.size)
    if not sizes:
        return None
    sizes.sort()
    return sizes[len(sizes) // 2]


def _is_all_bold(para) -> bool:
    """Check if all non-empty runs in a paragraph are bold."""
    runs_with_text = [r for r in para.runs if r.text.strip()]
    if not runs_with_text:
        return False
    return all(r.bold for r in runs_with_text)


def _has_large_font(para, median_size) -> bool:
    """Check if any run has font size significantly larger than median."""
    if median_size is None:
        return False
    threshold = median_size + Pt(2)
    for run in para.runs:
        if run.font.size and run.font.size > threshold:
            return True
    return False


def _is_all_caps_text(text: str) -> bool:
    """Check if text is ALL CAPS (and has letters)."""
    letters = [c for c in text if c.isalpha()]
    return len(letters) >= 2 and all(c.isupper() for c in letters)


def _looks_like_sub_entry(text: str) -> bool:
    """Check if text looks like a sub-entry (has dates, pipes, dashes with years)."""
    # "Company Name | 2019-2022" or "Ben-Gurion University | 2023-Present"
    if re.search(r"\|\s*\d{4}", text):
        return True
    # "Some Role, Organization | YYYY-YYYY"
    if re.search(r"\d{4}\s*[-–]\s*(\d{4}|[Pp]resent)", text):
        return True
    return False


def is_section_header(para, median_font_size, para_index: int = -1) -> bool:
    """Determine if a paragraph is a section header using three signals.

    Args:
        para: The paragraph to check.
        median_font_size: Median font size across the document.
        para_index: Index of this paragraph (-1 if unknown).
    """
    text = para.text.strip()
    if not text:
        return False

    word_count = len(text.split())

    # Never treat the first paragraph as a section header — it's the name
    if para_index == 0:
        return False

    # Sub-entries with dates are not headers (e.g. "Company | 2019-2022")
    if _looks_like_sub_entry(text):
        return False

    # Signal 1: heading style
    style_name = para.style.name.lower() if para.style and para.style.name else ""
    if "heading" in style_name or style_name == "title":
        return True

    # Signal 2: formatting-based (short + bold/large/caps)
    if word_count <= 5:
        if _is_all_bold(para):
            return True
        if _has_large_font(para, median_font_size):
            return True
        if _is_all_caps_text(text):
            return True

    # Signal 3: content-based (known section name)
    if text.lower() in _KNOWN_NAMES:
        return True

    return False


# ---------------------------------------------------------------------------
# Contact block detection
# ---------------------------------------------------------------------------

def _is_contact_paragraph(para) -> bool:
    """Check if a paragraph contains contact info (email, phone, linkedin)."""
    text = para.text
    return bool(
        _EMAIL_RE.search(text)
        or _LINKEDIN_RE.search(text)
        or (len(text.strip()) < 100 and _PHONE_RE.search(text))
    )


# ---------------------------------------------------------------------------
# Main section map builder
# ---------------------------------------------------------------------------

def build_section_map(doc: Document) -> DocxSectionMap:
    """Analyze a docx Document and return a map of detected sections."""
    paragraphs = doc.paragraphs
    if not paragraphs:
        return DocxSectionMap(header=None, contact=None)

    median_font_size = _compute_median_font_size(paragraphs)

    # Pass 1: find all section headers
    headers = []
    for i, para in enumerate(paragraphs):
        if is_section_header(para, median_font_size, para_index=i):
            headers.append((i, para.text.strip()))

    # Handle multi-line headers: merge consecutive short bold paragraphs
    merged_headers = []
    i = 0
    while i < len(headers):
        idx, text = headers[i]
        # Check if next header is immediately after and also short
        if (i + 1 < len(headers)
                and headers[i + 1][0] == idx + 1
                and len(text.split()) <= 3
                and len(headers[i + 1][1].split()) <= 3):
            merged_text = f"{text} {headers[i + 1][1]}"
            merged_headers.append((idx, merged_text))
            i += 2  # skip the next one
        else:
            merged_headers.append((idx, text))
            i += 1
    headers = merged_headers

    # Identify the "header" block (name/title) — everything before first section header
    first_header_idx = headers[0][0] if headers else len(paragraphs)

    # Filter out name/title paragraphs from being treated as section headers
    # The first 1-3 paragraphs before any section header are likely name/title/contact
    # But only if they were detected as headers due to formatting (not content match)
    real_headers = []
    for idx, text in headers:
        if idx < first_header_idx:
            # This is in the header block — skip it as a section header
            continue
        # Also skip if it's actually the person's name (first non-empty paragraph)
        real_headers.append((idx, text))

    # Rebuild first_header_idx with real headers
    first_header_idx = real_headers[0][0] if real_headers else len(paragraphs)

    # Build header section (name + title, before first real section)
    header_section = None
    if first_header_idx > 0:
        header_section = DocxSection(
            name="header",
            semantic_type="header",
            header_para=-1,
            start_para=0,
            end_para=first_header_idx - 1,
        )

    # Detect contact block within the header area
    contact_section = None
    for i in range(min(first_header_idx, len(paragraphs))):
        if _is_contact_paragraph(paragraphs[i]):
            contact_section = DocxSection(
                name="contact",
                semantic_type="contact",
                header_para=-1,
                start_para=i,
                end_para=i,  # may span multiple — extend below
            )
            # Extend contact block for consecutive contact paragraphs
            for j in range(i + 1, min(first_header_idx, len(paragraphs))):
                if _is_contact_paragraph(paragraphs[j]):
                    contact_section.end_para = j
                else:
                    break
            break

    # Build content sections from real headers
    sections = []
    for j, (header_idx, header_text) in enumerate(real_headers):
        next_idx = real_headers[j + 1][0] if j + 1 < len(real_headers) else len(paragraphs)
        sections.append(DocxSection(
            name=header_text,
            semantic_type=classify_section(header_text),
            header_para=header_idx,
            start_para=header_idx + 1,
            end_para=next_idx - 1,
        ))

    # Find unmapped paragraphs
    mapped = set()
    if header_section:
        mapped.update(range(header_section.start_para, header_section.end_para + 1))
    for s in sections:
        mapped.add(s.header_para)
        mapped.update(range(s.start_para, s.end_para + 1))
    unmapped = [i for i in range(len(paragraphs)) if i not in mapped]

    return DocxSectionMap(
        header=header_section,
        contact=contact_section,
        sections=sections,
        unmapped_paras=unmapped,
    )


def parse_docx(path: Path) -> tuple[Document, DocxSectionMap]:
    """Open a docx file and return the Document + its section map."""
    doc = Document(str(path))
    section_map = build_section_map(doc)
    return doc, section_map


def extract_section_text(doc: Document, section: DocxSection) -> str:
    """Extract plain text from a section's paragraph range."""
    paragraphs = doc.paragraphs
    lines = []
    start = section.start_para
    end = min(section.end_para, len(paragraphs) - 1)
    for i in range(start, end + 1):
        text = paragraphs[i].text.strip()
        if text:
            lines.append(text)
    return "\n".join(lines)


def extract_all_text(doc: Document, section_map: DocxSectionMap) -> dict:
    """Extract text from all sections as a dict keyed by section name."""
    result = {}
    if section_map.header:
        result["header"] = extract_section_text(doc, section_map.header)
    if section_map.contact:
        result["contact"] = extract_section_text(doc, section_map.contact)
    for s in section_map.sections:
        result[s.name] = extract_section_text(doc, s)
    return result
